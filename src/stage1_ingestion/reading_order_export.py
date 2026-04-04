"""Export page_blocks to PDF, DOCX, or Markdown in reading order; respects <headingN> tags from preserve-heading."""

from __future__ import annotations

import html
import re
from pathlib import Path
from typing import Any, Dict, List, Tuple

import fitz

_TAG_RE = re.compile(
    r"<(heading1|heading2|heading3|body|caption|comment)>(.*?)</\1>",
    re.DOTALL,
)

def _sort_key_row(r: Dict[str, Any]) -> Tuple[int, float, float]:
    bbox = r.get("bbox") or [0, 0, 0, 0]
    if not isinstance(bbox, list) or len(bbox) < 4:
        return (int(r.get("page", 0) or 0), 0.0, 0.0)
    return (int(r.get("page", 0) or 0), float(bbox[1]), float(bbox[0]))


def sort_blocks_reading_order(rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    return sorted(rows, key=_sort_key_row)


def _estimate_font_size_for_bbox(page: fitz.Page, bbox: Tuple[float, float, float, float]) -> float:
    x0, y0, x1, y1 = bbox
    sizes: List[float] = []
    try:
        td = page.get_text("dict")
        for block in td.get("blocks") or []:
            if block.get("type") != 0:
                continue
            for line in block.get("lines") or []:
                for span in line.get("spans") or []:
                    sb = span.get("bbox")
                    if not sb or len(sb) < 4:
                        continue
                    sx0, sy0, sx1, sy1 = float(sb[0]), float(sb[1]), float(sb[2]), float(sb[3])
                    if sx1 <= x0 or sx0 >= x1 or sy1 <= y0 or sy0 >= y1:
                        continue
                    try:
                        sizes.append(float(span.get("size", 11)))
                    except (TypeError, ValueError):
                        sizes.append(11.0)
    except (RuntimeError, ValueError, AttributeError):
        pass
    if sizes:
        return max(6.0, min(24.0, sum(sizes) / len(sizes)))
    return 11.0


def _strip_sup_sub(s: str) -> str:
    return re.sub(r"</?(?:sup|sub)>", "", s)


def parse_heading_segments(raw: str) -> List[Tuple[str, str]]:
    """Split raw_text into (role, inner_text) using <heading1>...</heading1> etc.; HTML-unescape inner."""
    if not raw or not raw.strip():
        return []
    s = raw.strip()
    if "<heading1" not in s and "<heading2" not in s and "<body" not in s and "<caption" not in s:
        return [("body", _strip_sup_sub(s))]
    parts: List[Tuple[str, str]] = []
    pos = 0
    for m in _TAG_RE.finditer(s):
        role = m.group(1)
        inner = m.group(2)
        parts.append((role, html.unescape(_strip_sup_sub(inner))))
    if parts:
        return parts
    return [("body", _strip_sup_sub(s))]


def _role_to_pdf_font(role: str) -> Tuple[str, float, bool]:
    if role == "heading1":
        return ("Helvetica-Bold", 18.0, True)
    if role == "heading2":
        return ("Helvetica-Bold", 14.0, True)
    if role == "heading3":
        return ("Helvetica-Bold", 12.0, True)
    if role == "caption":
        return ("Helvetica-Oblique", 9.0, False)
    if role == "comment":
        return ("Helvetica", 8.0, False)
    return ("Helvetica", 11.0, False)


def _table_raw_to_md_lines(raw: str) -> str:
    lines = [ln for ln in raw.splitlines() if ln.strip()]
    if not lines:
        return ""
    rows_cells = [row.split("\t") for row in lines]
    ncol = max(len(r) for r in rows_cells)
    norm = [r + [""] * (ncol - len(r)) for r in rows_cells]
    header = norm[0]
    sep = "|" + "|".join("---" for _ in header) + "|"
    top = "|" + "|".join(c.strip() for c in header) + "|"
    out = [top, sep]
    for r in norm[1:]:
        out.append("|" + "|".join(c.strip() for c in r) + "|")
    return "\n".join(out) + "\n\n"


def export_page_blocks_to_pdf(
    rows: List[Dict[str, Any]],
    source_pdf: Path,
    out_path: Path,
) -> None:
    try:
        from reportlab.pdfgen import canvas as rl_canvas
    except ImportError as e:
        raise RuntimeError("export_page_blocks_to_pdf requires reportlab: pip install reportlab") from e

    rows = sort_blocks_reading_order(rows)
    doc = fitz.open(source_pdf)
    try:
        c = rl_canvas.Canvas(str(out_path))
        for pi in range(doc.page_count):
            page = doc[pi]
            pw = float(page.rect.width)
            ph = float(page.rect.height)
            c.setPageSize((pw, ph))
            page_num = pi + 1
            page_rows = [r for r in rows if int(r.get("page", 0) or 0) == page_num]
            for r in page_rows:
                raw = (r.get("raw_text") or "").strip()
                bbox = r.get("bbox")
                if not isinstance(bbox, list) or len(bbox) < 4:
                    continue
                x0, y0, x1, y1 = float(bbox[0]), float(bbox[1]), float(bbox[2]), float(bbox[3])
                bt = r.get("block_type") or "text"

                if bt == "image":
                    c.setStrokeColorRGB(0.5, 0.5, 0.5)
                    c.rect(x0, ph - y1, x1 - x0, y1 - y0)
                    c.stroke()
                    c.setFont("Helvetica", 9)
                    label = (raw or "[image]").splitlines()[0][:200]
                    c.drawString(x0 + 2, ph - y0 - 12, label)
                    continue

                if bt == "table":
                    fs = max(8.0, _estimate_font_size_for_bbox(page, (x0, y0, x1, y1)) * 0.9)
                    c.setFont("Helvetica", fs)
                    plain = _strip_sup_sub(raw)
                    line_h = max(fs * 1.2, 9.0)
                    y_base = ph - y0 - line_h
                    for li, line in enumerate(plain.splitlines()[:60]):
                        c.drawString(x0, y_base - li * line_h, line[:1200])
                    continue

                if bt == "equation":
                    c.setFont("Courier", 10)
                    plain = _strip_sup_sub(raw)
                    line_h = 12.0
                    y_base = ph - y0 - line_h
                    for li, line in enumerate(plain.splitlines()[:40]):
                        c.drawString(x0, y_base - li * line_h, line[:1200])
                    continue

                if not raw:
                    continue

                segs = parse_heading_segments(raw)
                y_cursor = ph - y0 - 12.0
                for role, chunk in segs:
                    if not chunk.strip():
                        continue
                    font_name, fs, _ = _role_to_pdf_font(role)
                    c.setFont(font_name, fs)
                    line_h = max(fs * 1.2, 9.0)
                    for li, line in enumerate(chunk.splitlines()[:40]):
                        c.drawString(x0, y_cursor - li * line_h, line[:1200])
                    y_cursor -= line_h * max(1, len(chunk.splitlines()[:40])) + 4
            c.showPage()
        c.save()
    finally:
        doc.close()


def export_page_blocks_to_docx(
    rows: List[Dict[str, Any]],
    source_pdf: Path,
    out_path: Path,
) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as e:
        raise RuntimeError("export_page_blocks_to_docx requires python-docx: pip install python-docx") from e

    rows = sort_blocks_reading_order(rows)
    docx = Document()
    pdf_doc = fitz.open(source_pdf)
    try:
        page_cache: Dict[int, fitz.Page] = {i + 1: pdf_doc[i] for i in range(pdf_doc.page_count)}
        for r in rows:
            raw = (r.get("raw_text") or "").strip()
            bbox = r.get("bbox")
            pnum = int(r.get("page", 0) or 0)
            page = page_cache.get(pnum)
            fs_base = 11.0
            if page and isinstance(bbox, list) and len(bbox) >= 4:
                fs_base = _estimate_font_size_for_bbox(
                    page, (float(bbox[0]), float(bbox[1]), float(bbox[2]), float(bbox[3]))
                )

            bt = r.get("block_type") or "text"

            if bt == "image":
                p = docx.add_paragraph()
                p.add_run("[IMAGE]").bold = True
                if raw:
                    p.add_run("\n" + _strip_sup_sub(raw))
                continue

            if bt == "table":
                p = docx.add_paragraph()
                run = p.add_run(_strip_sup_sub(raw))
                run.font.size = Pt(min(24.0, max(6.0, fs_base)))
                continue

            if bt == "equation":
                p = docx.add_paragraph()
                run = p.add_run(_strip_sup_sub(raw))
                run.font.name = "Courier New"
                run.font.size = Pt(10)
                continue

            if not raw:
                continue

            segs = parse_heading_segments(raw)
            for role, chunk in segs:
                if not chunk.strip():
                    continue
                para = docx.add_paragraph()
                if role == "heading1":
                    para.style = "Heading 1"
                    para.add_run(chunk)
                elif role == "heading2":
                    para.style = "Heading 2"
                    para.add_run(chunk)
                elif role == "heading3":
                    para.style = "Heading 3"
                    para.add_run(chunk)
                else:
                    run = para.add_run(chunk)
                    run.font.size = Pt(min(24.0, max(6.0, fs_base * (0.85 if role in ("caption", "comment") else 1.0))))
                    if role == "caption":
                        run.italic = True
        docx.save(str(out_path))
    finally:
        pdf_doc.close()


def export_page_blocks_to_markdown(
    rows: List[Dict[str, Any]],
    source_pdf: Path,
    out_path: Path,
) -> None:
    del source_pdf  # page breaks optional; not required to match PDF
    rows = sort_blocks_reading_order(rows)
    lines: List[str] = []
    cur_page = 0
    for r in rows:
        pnum = int(r.get("page", 0) or 0)
        if pnum != cur_page:
            if lines:
                lines.append("")
            lines.append(f"<!-- page {pnum} -->")
            lines.append("")
            cur_page = pnum

        raw = (r.get("raw_text") or "").strip()
        bt = r.get("block_type") or "text"

        if bt == "image":
            lines.append("```")
            lines.append("[IMAGE BLOCK]")
            if raw:
                lines.append(raw)
            lines.append("```")
            lines.append("")
            continue

        if bt == "table":
            lines.append(_table_raw_to_md_lines(raw))
            continue

        if bt == "equation":
            lines.append("```")
            lines.append(raw)
            lines.append("```")
            lines.append("")
            continue

        if not raw:
            continue

        segs = parse_heading_segments(raw)
        for role, chunk in segs:
            if not chunk.strip():
                continue
            if role == "heading1":
                lines.append("# " + chunk.replace("\n", " ").strip())
            elif role == "heading2":
                lines.append("## " + chunk.replace("\n", " ").strip())
            elif role == "heading3":
                lines.append("### " + chunk.replace("\n", " ").strip())
            elif role == "caption":
                lines.append(f"*{chunk.strip()}*")
            elif role == "comment":
                lines.append(f"<!-- {chunk.strip()} -->")
            else:
                lines.append(chunk.strip())
            lines.append("")
    out_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")
